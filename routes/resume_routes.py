from fastapi import APIRouter, HTTPException, Request, Form
from fastapi.templating import Jinja2Templates
from fastapi.responses import RedirectResponse, FileResponse
from services.resume_service import generate_resume, save_resume, get_all_resumes, view_resume
from models.resume_model import Resume
from pdfcrowd import HtmlToPdfClient
import os

router = APIRouter()
templates = Jinja2Templates(directory="templates")

@router.get("/create_resume")
async def show_form(request: Request):
    return templates.TemplateResponse("business_problem_form.html", {"request": request})

@router.post("/generate_response")
async def create_resume(client_problem: str = Form(...)):
    generated_resume = generate_resume(client_problem)
    
    # Create a Resume instance
    resume = Resume(client_problem=client_problem, generated_resume=generated_resume)
    
    """# Convert the Pydantic model instance to a dictionary
    resume_dict = resume.dict()
    
    # Save the resume to the database
    resume_id = save_resume(resume_dict)"""

    return {"generated_resume": resume}
    
"""    if resume_id:
        # Redirect to the resume list page
        return RedirectResponse(url="/", status_code=303)
    else:
        raise HTTPException(status_code=500, detail="Failed to save resume")"""

@router.get("/")
async def list_resumes(request: Request):
    resumes = get_all_resumes()
    return templates.TemplateResponse("resume_list.html", {"request": request, "resumes": resumes})

@router.get("/create/")
async def create(request: Request):
    resumes = get_all_resumes()
    return templates.TemplateResponse("riceffile.html", {"request": request, "resumes": resumes})

"""@router.get("/resume_view/{resume_name}")
async def view(resume_name: str, request: Request):
    # Call the synchronous function to retrieve the resume data
    print(resume_name)
    resume = view_resume(resume_name)
    print(resume)
    if isinstance(resume, list):
        print("if block of view excuting")
        return templates.TemplateResponse("riceffile.html", {"request": request, "resume": resume})
    else:
        raise HTTPException(status_code=404, detail="RICEF not found")"""



# new code 
from fastapi import APIRouter, HTTPException, Request, Response
from fastapi.responses import FileResponse
from docxtpl import DocxTemplate
import io


"""
@router.get("/resume_view/{resume_name}")
async def view(resume_name: str, request: Request):
    # Retrieve the resume data
    resume = view_resume(resume_name)
    if isinstance(resume, list):
        # Load Word template and populate with data
        template = DocxTemplate("templates/template.docx")
        template.render({"resume": resume})

        # Save to a BytesIO stream instead of file
        byte_io = io.BytesIO()
        template.save(byte_io)
        byte_io.seek(0)

        # Return the Word document as a downloadable file
        return Response(byte_io.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
            "Content-Disposition": f"attachment; filename={resume_name}.docx"
        })
    else:
        raise HTTPException(status_code=404, detail="RICEF not found")
"""


"""@router.get("/resume_download/{resume_name}")
async def download_pdf(resume_name: str, request: Request):
    # Retrieve the resume data
    resume = view_resume(resume_name)
    filename = resume_name
    print("Printing name of the file")
    print(filename)
    print("After printing file name")
    if isinstance(resume, list):
        # Load Word template and populate with data
        template = DocxTemplate("templates/{resume_name}.docx")
        template.render({"resume": resume})

        # Save to a BytesIO stream instead of file
        byte_io = io.BytesIO()
        template.save(byte_io)
        byte_io.seek(0)

        # Return the Word document as a downloadable file
        return Response(byte_io.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
            "Content-Disposition": f"attachment; filename={resume_name}.docx"
        })
    else:
        raise HTTPException(status_code=404, detail="Resume not found")
    """

"""
@router.get("/resume_download/{resume_name}")
async def download_docx(resume_name: str):
    import os
    from fastapi.responses import FileResponse
    from fastapi import HTTPException

    # Define the path to the specific template file
    file_path = f"templates/{resume_name}.docx"
    fill_values(file_path)
    
    # Check if the specified file exists in the templates folder
    if not os.path.isfile(file_path):
        raise HTTPException(status_code=404, detail=f"File '{resume_name}.docx' not found in templates folder.")

    # Serve the file as a downloadable attachment
    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"new_{resume_name}.docx"
    )
"""

from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
import shutil



# Path to your template file
TEMPLATE_PATH = "templates/template.docx"
# Path to save the updated document
OUTPUT_PATH = "templates/dynamic_sections_list.docx"

# Load the document template when the server starts
shutil.copy(TEMPLATE_PATH, OUTPUT_PATH)  # Copy the template to the output path
doc = Document(OUTPUT_PATH)

# List to store sections and subsections as they are added
sections_data = []

class SectionRequest(BaseModel):
    title: str
    content: str
    type: str 
    #user_name: str
    select_content: str
    customerName: str
    customerShortName: str
    aeName: str
    titleRfpTemplate: str
    

def fill_values(file_path, request):
    tpl = DocxTemplate(file_path)
    context = {
        #"user_name": request.user_name,
        "title": request.title,
        "content": request.content,
        "type": request.type,
        "select_content": request.select_content,
        "customerName": request.customerName,
        "customerShortName": request.customerShortName,
        "aeName": request.aeName,
        "titleRfpTemplate": request.titleRfpTemplate,
    }
    tpl.render(context)
    tpl.save(file_path)


@router.get("/resume_download/{resume_name}")
async def download_docx(resume_name: str):
    file_path = f"templates/{resume_name}.docx"
    
    # Check if the specified file exists in the templates folder
    if not os.path.isfile(file_path):
        raise HTTPException(status_code=404, detail=f"File '{resume_name}.docx' not found in templates folder.")
    
    # Fill the values before returning the file
    #fill_values(file_path, request)
    
    # Serve the file as a downloadable attachment
    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"new_{resume_name}.docx"
    )

def add_sections_from_list(doc, sections):
    """
    Adds sections and subsections to a Word document dynamically from a list of dictionaries.
    """
    for section in sections:
        # Add main section heading
        doc.add_heading(section['title'], level=2)
        if 'content' in section:
            doc.add_paragraph(section['content'])

        # Check if there are subsections
        if 'subsections' in section:
            for subsection in section['subsections']:
                # Add subsection heading
                doc.add_heading(subsection['title'], level=3)
                if 'content' in subsection:
                    doc.add_paragraph(subsection['content'])
"""@router.post("/add_section")
async def add_section(Request: SectionRequest):
    # Define the path for the project file and template
    file_path = f"templates/{Request.user_name}.docx"
    template_path = "templates/template.docx"
    
    
    # Check if the project file already exists
    if os.path.exists(file_path):
        doc = Document(file_path)  # Open existing document
    else:
        # Render the template with the user_name placeholder
        tpl = DocxTemplate(template_path)
        context = {"user_name": Request.user_name}
        tpl.render(context)
        tpl.save(file_path)  # Save the rendered document

        doc = Document(file_path)  # Open the newly rendered document
    # Add content based on the type specified

    source = "content_word_ducuments/NTT DATA Business Solutions Overview.docx"
    if Request.type == "section":
        doc.add_heading(Request.title, level=2)
        doc.add_paragraph(Request.content)
    elif Request.type == "subsection":
        doc.add_heading(Request.title, level=3)
        doc.add_paragraph(Request.content)

    # Save the document
    doc.save(file_path)

    return JSONResponse(content={"message": "Section or Subsection added successfully."})"""




from fastapi import APIRouter
from fastapi.responses import JSONResponse
from docx import Document
from docxcompose.composer import Composer
import os
"""
@router.post("/add_section")
async def add_section(request: SectionRequest):
    # Define paths
    file_path = f"templates/{request.user_name}.docx"
    template_path = "templates/template.docx"
    
    source_file = f"{request.select_content}.docx"
    # Resolve the absolute path for the source document
    base_path = os.path.dirname(os.path.abspath(__file__))  # Current script's directory
    root_path = os.path.dirname(base_path)  # Parent directory of the script (root of the project)
    source = os.path.join(root_path, "content_word_documents", source_file)

    # Check if the document exists
    if os.path.exists(file_path):
        doc = Document(file_path)  # Open existing document
    else:
        # Render the template with the user_name placeholder
        tpl = DocxTemplate(template_path)
        context = {"user_name": request.user_name}
        tpl.render(context)
        tpl.save(file_path)
        doc = Document(file_path)  # Open the newly created document

    # Add section or subsection
    if request.type == "section":
        doc.add_heading(request.title, level=2)
    elif request.type == "subsection":
        doc.add_heading(request.title, level=3)
    else:
        return JSONResponse(content={"message": "Invalid type specified."}, status_code=400)

    doc.add_paragraph(request.content)

    # Save current updates to the document
    doc.save(file_path)

    # Append the entire source document
    destination_doc = Document(file_path)
    
    # Check if the source document exists
    if not os.path.exists(source):
        return JSONResponse(content={"message": "Source document not found."}, status_code=404)

    source_doc = Document(source)

    # Use Composer to merge documents
    composer = Composer(destination_doc)
    composer.append(source_doc)

    # Save the updated document with merged content
    composer.save(file_path)

    return JSONResponse(content={"message": "Section, Subsection, and source content appended successfully."})
"""


"""
@router.post("/add_section")
async def add_section(request: SectionRequest):
    # Define paths
    file_path = f"templates/{request.user_name}.docx"
    template_path = "templates/template.docx"
    
    source_file = f"{request.select_content}.docx"
    # Resolve the absolute path for the source document
    base_path = os.path.dirname(os.path.abspath(__file__))  # Current script's directory
    root_path = os.path.dirname(base_path)  # Parent directory of the script (root of the project)
    source = os.path.join(root_path, "content_word_documents", source_file)

    # Initialize section and subsection counters
    section_counter = 1
    subsection_counter = 1

    # Check if the document exists
    if os.path.exists(file_path):
        doc = Document(file_path)  # Open existing document
    else:
        # If the file doesn't exist, create it from template
        tpl = DocxTemplate(template_path)
        context = {"user_name": request.user_name,
                   "customerName": request.customerName,}  # Context to render the placeholder
        tpl.render(context)
        tpl.save(file_path)  # Save the template as the initial file
        doc = Document(file_path)  # Open the newly created document

    # Add section or subsection
    if request.type == "section":
        # Add section heading (e.g., 1)
        doc.add_heading(f"{request.title}", level=1)
        section_counter += 1  # Increment section counter for next section
    elif request.type == "subsection":
        # Add subsection heading (e.g., 1.1)
        doc.add_heading(f"{request.title}", level=2)
        subsection_counter += 1  # Increment subsection counter for next subsection
    else:
        return JSONResponse(content={"message": "Invalid type specified."}, status_code=400)

    # Add the content under the section/subsection
    doc.add_paragraph(request.content)

    # Save current updates to the document
    doc.save(file_path)

    # Append the entire source document
    destination_doc = Document(file_path)
    
    # Check if the source document exists
    if not os.path.exists(source):
        return JSONResponse(content={"message": "Source document not found."}, status_code=404)

    source_doc = Document(source)

    # Use Composer to merge documents
    composer = Composer(destination_doc)
    composer.append(source_doc)

    # Save the updated document with merged content
    composer.save(file_path)

    tpl = DocxTemplate(file_path)
    context = {
        "user_name": request.user_name,
        "title": request.title,
        "content": request.content,
        "type": request.type,
        "select_content": request.select_content,
        "customerName": request.customerName,
        "customerShortName": request.customerShortName,
        "aeName": request.aeName,
        "titleRfpTemplate": request.titleRfpTemplate,
    }
    tpl.render(context)
    tpl.save(file_path)


    return JSONResponse(content={"message": "Section, Subsection, and source content appended successfully."})
"""








"""
@router.post("/add_section")
async def add_section(request: SectionRequest):
    # Define paths
    file_path = f"templates/{request.user_name}.docx"
    template_path = "templates/template.docx"
    source_file = f"{request.select_content}.docx"
    
    # Resolve the absolute path for the source document
    base_path = os.path.dirname(os.path.abspath(__file__))  # Current script's directory
    root_path = os.path.dirname(base_path)  # Parent directory of the script (root of the project)
    source = os.path.join(root_path, "content_word_documents", source_file)

    # Define the context for rendering placeholders
    context = {
        "user_name": request.user_name,
        "title": request.title,
        "content": request.content,
        "type": request.type,
        "select_content": request.select_content,
        "customerName": request.customerName,
        "customerShortName": request.customerShortName,
        "aeName": request.aeName,
        "titleRfpTemplate": request.titleRfpTemplate,
    }

    # Initialize the document
    if os.path.exists(file_path):
        # Open an existing document
        doc = Document(file_path)
    else:
        # Create a new document by rendering the template
        tpl = DocxTemplate(template_path)
        tpl.render(context)  # Render placeholders
        tpl.save(file_path)  # Save the rendered template as the user's document
        doc = Document(file_path)

    # Render placeholders in the source document (if required)
    if os.path.exists(source):
        source_tpl = DocxTemplate(source)
        source_tpl.render(context)  # Render placeholders in the source document
        source_tpl.save(source)  # Save the rendered source document
    else:
        return JSONResponse(content={"message": "Source document not found."}, status_code=404)

    # Add section or subsection
    if request.type == "section":
        doc.add_heading(f"{request.title}", level=1)
    elif request.type == "subsection":
        doc.add_heading(f"{request.title}", level=2)
    else:
        return JSONResponse(content={"message": "Invalid type specified."}, status_code=400)

    # Add the content under the section/subsection
    doc.add_paragraph(request.content)
    doc.save(file_path)  # Save the updated document

    # Append the source document to the user document
    destination_doc = Document(file_path)
    source_doc = Document(source)
    composer = Composer(destination_doc)
    composer.append(source_doc)
    composer.save(file_path)  # Save the merged document

    return JSONResponse(content={"message": "Section, subsection, and source content appended successfully."})
"""


@router.post("/add_section")
async def add_section(request: SectionRequest):
    # Define paths
    file_path = f"templates/{request.titleRfpTemplate}.docx"
    template_path = "templates/template.docx"
    source_file = f"{request.select_content}.docx"
    
    # Resolve paths
    base_path = os.path.dirname(os.path.abspath(__file__))  # Current script's directory
    root_path = os.path.dirname(base_path)  # Parent directory of the script
    source = os.path.join(root_path, "content_word_documents", source_file)
    rendered_source_path = os.path.join(root_path, "content_word_documents", "rendered_documents", source_file)
    
    # Ensure the rendered documents folder exists
    os.makedirs(os.path.dirname(rendered_source_path), exist_ok=True)

    # Define context for rendering placeholders
    context = {
        #"user_name": request.user_name,
        "title": request.title,
        "content": request.content,
        "type": request.type,
        "select_content": request.select_content,
        "customerName": request.customerName,
        "customerShortName": request.customerShortName,
        "aeName": request.aeName,
        "titleRfpTemplate": request.titleRfpTemplate,
    }

    # Initialize the user document
    if os.path.exists(file_path):
        # Open an existing document
        doc = Document(file_path)
    else:
        # Create a new document from the template
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(file_path)
        doc = Document(file_path)

    # Render placeholders in the source document and save as a new rendered file
    if os.path.exists(source):
        source_tpl = DocxTemplate(source)
        source_tpl.render(context)  # Render placeholders
        source_tpl.save(rendered_source_path)  # Save the rendered source document
    else:
        return JSONResponse(content={"message": "Source document not found."}, status_code=404)

    # Add section or subsection
    if request.type == "section":
        doc.add_heading(f"{request.title}", level=1)
    elif request.type == "subsection":
        doc.add_heading(f"{request.title}", level=2)
    else:
        return JSONResponse(content={"message": "Invalid type specified."}, status_code=400)

    # Add the content under the section/subsection
    doc.add_paragraph(request.content)
    doc.save(file_path)  # Save the updated document

    # Append the rendered source document to the user document
    destination_doc = Document(file_path)
    rendered_source_doc = Document(rendered_source_path)
    composer = Composer(destination_doc)
    composer.append(rendered_source_doc)
    composer.save(file_path)  # Save the merged document

    return JSONResponse(content={"message": "Section, subsection, and source content appended successfully."})
